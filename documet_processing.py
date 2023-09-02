import re
import os
import pypandoc
import string
import shutil
# from pprint import pprint
from docx import Document

import constant



SECTION_PATTERN_MULTIPLE_LINES = re.compile("|".join([
    constant.CHUONG_PATTERN_MULTIPLE_LINES,
    constant.PHULUC_PATTERN_MULTIPLE_LINES,
]))

SECTION_PATTERN = re.compile("|".join([
    constant.CHUONG_PATTERN,
    constant.PHULUC_PATTERN,
]))


def get_name(file_name):
    return file_name.split('/')[-1].split('.')[0]
 
# =================================================================================
# 1. REMOVE TABLES
# =================================================================================

doc_file_names = [
        (constant.POLICY_WORD_FOLDER + '/' + text)
        for text in os.listdir(constant.POLICY_WORD_FOLDER)
    ]

def delete_table(word_file):
    document = Document(word_file)
    processed_file_name = constant.POLICY_PROCESSED_WORD_FOLDER + '/' + get_name(word_file) + ".docx"
    number_of_table = len(document.tables)
    for i in range(number_of_table):
        document.tables[0]._element.getparent().remove(document.tables[0]._element)
    document.save(processed_file_name)

    return processed_file_name

processed_doc_file_names = list(map(delete_table, doc_file_names))

# =================================================================================
# 2. CONVERT MS WORD FILES TO TEXT FILES
# =================================================================================

def convert_word_to_txt(processed_word_file):
    pypandoc.download_pandoc()
    text_file_name = constant.POLICY_TEXT_FOLDER + '/' + get_name(processed_word_file) + ".txt"
    pypandoc.convert_file(processed_word_file, 'plain', outputfile=text_file_name)
    return text_file_name


text_file_names = list(map(convert_word_to_txt, processed_doc_file_names))

# text_file_names = [
#         (constant.POLICY_TEXT_FOLDER + '/' + text)
#         for text in os.listdir(constant.POLICY_TEXT_FOLDER)
#     ]

# =================================================================================
# 3. READ TEXT FILES CONTENTS
# =================================================================================

def read_lines(text_file_name):
    with open(text_file_name) as text_file:
        lines = text_file.readlines()
    ret = list()
    cur_line = ""
    for line in lines:
        if not line.strip():
            ret.append(cur_line)
            ret.append("\n")
            cur_line = ""
            continue
        
        if not cur_line:
            cur_line = line.strip()
        else:
            cur_line += " "
            cur_line += line.strip()
    ret.append(cur_line)
    return ret

content_list = list(map(read_lines, text_file_names))

def handle_multiple_lines_section(lines):
    index = 0
    new_lines = list()
    while True:
        if index == len(lines)-1:
            new_lines.append(lines[index])
            return new_lines 
        if re.match(SECTION_PATTERN_MULTIPLE_LINES, lines[index]):
            new_lines.append(lines[index] + " " + lines[index+2])
            index += 3
        else:
            new_lines.append(lines[index])
            index += 1

content_list = list(map(handle_multiple_lines_section, content_list))

# =================================================================================
# 4. GENERATE INDEX
# =================================================================================

def get_title(lines):
    for index, line in enumerate(lines):
            if constant.TITLE_FLAG in line.lower():
                return lines[index-2]
            
title_list = list(map(get_title, content_list))

def get_section(lines):    
    section_list = list()
    for line in lines:
        if re.match(SECTION_PATTERN, line):
            section_list.append(line)
            continue
    return section_list

section_list = list(map(get_section, content_list))

def get_sub_section(lines):
    subsection = list()
    isStart_PHULUC = False

    for line in lines:
        if re.match(constant.PHULUC_PATTERN, line):
            isStart_PHULUC = True

        if re.match(constant.DIEU_PATTERN, line):
            subsection.append(line)

        if isStart_PHULUC and re.match(constant.PHU_LUC_SUB_SECTION_PATTERN, line):
            subsection.append(line)

    return subsection

subsection_list = list(map(get_sub_section, content_list))

# =================================================================================
# 5. GENERATE PASSAGE
# =================================================================================

def gen_passage(title, section, subsection, text, subsection_index):
    passage = ""
    passage += title
    passage += "\n\n"
    passage += section if section else "\n"
    passage += "\n"
    passage += subsection if subsection else "\n"
    passage += "\n\n"
    passage += text

    name = str(subsection_index) + '_' + title + '.txt'
    return name, passage

def get_text(title, section_list, subsection_list, lines):
    section_index = 0
    subsection_index = 0

    text_list = list()
    isSectionUpdate = False

    cur_section = ""
    cur_subsection = ""
    cur_text = ""

    for line in lines:
        if line == section_list[section_index]:
            isSectionUpdate = True
            continue

        if line == subsection_list[subsection_index]:
            if cur_subsection:                      
                text_list.append(gen_passage(title, cur_section, cur_subsection, cur_text, subsection_index))

            if isSectionUpdate:
                cur_section = section_list[section_index]
                section_index += 1 if section_index < len(section_list)-1 else 0
                isSectionUpdate = False
                
            cur_subsection = subsection_list[subsection_index]
            cur_text = ""
            subsection_index += 1 if subsection_index < len(subsection_list)-1 else 0
            continue
        
        if cur_subsection:
            cur_text += line    

    text_list.append(gen_passage(title, cur_section, cur_subsection, cur_text, subsection_index+1))

    return text_list

text_list = list(map(get_text, title_list, section_list, subsection_list, content_list))                

# =================================================================================
# 6. EXPORT PASSAGES
# =================================================================================
for index, title in enumerate(title_list):
    dir_path = constant.PASSAGE_FOLDER + '/' + title
    if os.path.exists(dir_path):
        shutil.rmtree(dir_path)


for index, title in enumerate(title_list):
    dir_path = constant.PASSAGE_FOLDER + '/' + title
    if not os.path.exists(dir_path):
        os.makedirs(constant.PASSAGE_FOLDER + '/' + title)
    for passage in text_list[index]:
        passage_name = passage[0]
        passage_content = passage[1]
        with open(dir_path + '/'+ passage_name, "w+") as file:
            file.write(passage_content)

# =================================================================================
# 7. CHECH SIZE LIMIT
# =================================================================================
def check_over_limit(file_name):
    with open(file_name, 'r') as f:
        return len(f.read().split()) > constant.PASSAGE_LENGTH_LIMIT


def split_file(file_name):
    if not check_over_limit(file_name):
        return
    
    with open(file_name, 'r') as f:
        lines = f.readlines()
    
    header = lines[:4]
    context = lines[6:]
    
    split_list = list()
    for line in context:
        if len(split_list) == 0:
             split_list.append(line)
        elif re.match(constant.SUBSUB_SECTION_PATERN, line):
            split_list.append(line)
        else:
            split_list[-1] += line

    for index, new_file in enumerate(split_list):
        folder_path = "/".join(file_name.split('/')[:-1])
        new_name = folder_path + "/" + file_name.split('/')[-1].split('_')[0] + list(string.ascii_lowercase)[index] + '_' + file_name.split('/')[-1].split('_')[1]
        new_content = "".join(header) + "\n\n" + new_file

        with open(new_name, "w+") as f:
            f.write(new_content)
    
    os.remove(file_name)

for folder_path in os.listdir(constant.PASSAGE_FOLDER):
    for filename in os.listdir(constant.PASSAGE_FOLDER + '/' + folder_path):
        split_file(constant.PASSAGE_FOLDER + '/' + folder_path+'/'+filename)